

import tkinter as tk
from tkinter import filedialog
import os
import fitz  # PyMuPDF
from collections import defaultdict
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload, MediaIoBaseDownload
from google.oauth2.credentials import Credentials

import io
import PyPDF2
from docx import Document
import time

SCOPES = ['https://www.googleapis.com/auth/drive.file']
MANUAL_DIR = 'manual_uploads'
os.makedirs(MANUAL_DIR, exist_ok=True)

# متغيرات لتخزين أزمنة آخر عمليات
last_search_time = None
last_sort_time = None
last_classify_time = None

def upload_to_drive(file_path):
    creds = Credentials.from_authorized_user_file('token.json', SCOPES)
    service = build('drive', 'v3', credentials=creds)
    
    folder_id = '1xIuZs_-AwFjM7D7Yc8-YWc9gXULl_SYG'  # معرف مجلد جوجل درايف الخاص بك
    
    file_metadata = {
        'name': os.path.basename(file_path),
        'parents': [folder_id]
    }
    media = MediaFileUpload(file_path, resumable=True)
    service.files().create(body=file_metadata, media_body=media).execute()
    print(f'✅ Uploaded to Drive: {file_path}')

def browse_and_save_manual_files():
    file_paths = filedialog.askopenfilenames(
        filetypes=[("Documents", "*.pdf *.docx")]
    )
    for path in file_paths:
        filename = os.path.basename(path)
        new_path = os.path.join(MANUAL_DIR, filename)
        with open(path, 'rb') as src, open(new_path, 'wb') as dst:
            dst.write(src.read())
        print(f'📁 Saved manually: {new_path}')
    status_label.config(text=f"✅ تم حفظ {len(file_paths)} ملف يدوياً.")

def upload_all_documents():
    all_files = []
    for filename in os.listdir(MANUAL_DIR):
        path = os.path.join(MANUAL_DIR, filename)
        if os.path.isfile(path):
            all_files.append(path)

    for file in all_files:
        upload_to_drive(file)

    status_label.config(text=f"✅ تم رفع {len(all_files)} ملف إلى Google Drive.")

def download_files_from_drive():
    creds = Credentials.from_authorized_user_file('token.json', SCOPES)
    service = build('drive', 'v3', credentials=creds)

    query = "mimeType='application/pdf' or mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'"

    results = service.files().list(q=query, fields="files(id, name)").execute()
    files = results.get('files', [])

    if not files:
        status_label.config(text="ℹ️ لم يتم العثور على ملفات PDF أو DOCX في Google Drive.")
        return

    for file in files:
        file_id = file['id']
        file_name = file['name']
        request = service.files().get_media(fileId=file_id)

        file_path = os.path.join(MANUAL_DIR, file_name)
        fh = open(file_path, 'wb')

        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while not done:
            status, done = downloader.next_chunk()

        fh.close()
        print(f'⬇️ تم تنزيل: {file_name}')

    status_label.config(text=f"✅ تم تنزيل {len(files)} ملف من Google Drive إلى {MANUAL_DIR}.")

def extract_title(file_path):
    if file_path.endswith('.pdf'):
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                if reader.metadata and reader.metadata.title:
                    return reader.metadata.title.strip()
                else:
                    first_page = reader.pages[0]
                    text = first_page.extract_text()
                    if text:
                        return text.strip().split('\n')[0]
        except Exception as e:
            print(f'❌ Error reading PDF: {file_path} | {e}')

    elif file_path.endswith('.docx'):
        try:
            doc = Document(file_path)
            for p in doc.paragraphs:
                if p.text.strip():
                    return p.text.strip()
        except Exception as e:
            print(f'❌ Error reading DOCX: {file_path} | {e}')

    return "عنوان غير معروف"

def get_documents_stats():
    total_size = 0
    total_files = 0
    for filename in os.listdir(MANUAL_DIR):
        file_path = os.path.join(MANUAL_DIR, filename)
        if os.path.isfile(file_path) and (filename.endswith('.pdf') or filename.endswith('.docx')):
            total_files += 1
            total_size += os.path.getsize(file_path)
    return total_files, total_size

def sort_documents_by_title():
    global last_sort_time

    start_time = time.time()

    documents = []

    for filename in os.listdir(MANUAL_DIR):
        file_path = os.path.join(MANUAL_DIR, filename)
        if os.path.isfile(file_path) and (filename.endswith('.pdf') or filename.endswith('.docx')):
            title = extract_title(file_path)
            documents.append((title.lower(), file_path))

    sorted_docs = sorted(documents, key=lambda x: x[0])

    elapsed = time.time() - start_time
    last_sort_time = elapsed

    result_window = tk.Toplevel(root)
    result_window.title("📄 المستندات المرتبة حسب العنوان")
    result_window.geometry("500x400")

    text_widget = tk.Text(result_window, wrap=tk.WORD)
    text_widget.pack(expand=True, fill=tk.BOTH)

    for title, path in sorted_docs:
        text_widget.insert(tk.END, f"• {title} → {os.path.basename(path)}\n")

    total_files, total_size = get_documents_stats()
    size_mb = total_size / (1024 * 1024)

    status_label.config(text=f"✅ تم فرز {len(sorted_docs)} ملف حسب العنوان. "
                             f"عدد المستندات: {total_files}، الحجم الكلي: {size_mb:.2f} ميجابايت، وقت الفرز: {elapsed:.2f} ثانية.")

def extract_text(file_path):
    text = ""
    try:
        if file_path.endswith('.pdf'):
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text() + "\n"
        elif file_path.endswith('.docx'):
            doc = Document(file_path)
            for p in doc.paragraphs:
                if p.text:
                    text += p.text + "\n"
    except Exception as e:
        print(f'❌ Error reading file: {file_path} | {e}')
    return text

def search_documents(keywords):
    if isinstance(keywords, str):
        keywords = [keywords.lower()]
    else:
        keywords = [kw.lower() for kw in keywords]

    matched_files = []

    for filename in os.listdir(MANUAL_DIR):
        if filename.endswith('.pdf') or filename.endswith('.docx'):
            file_path = os.path.join(MANUAL_DIR, filename)
            text = extract_text(file_path).lower()

            if all(kw in text for kw in keywords):
                matched_files.append((file_path, text))

    return matched_files

def highlight_text(text_widget, keywords):
    for kw in keywords:
        start_idx = "1.0"
        while True:
            pos = text_widget.search(kw, start_idx, nocase=1, stopindex=tk.END)
            if not pos:
                break
            end_pos = f"{pos}+{len(kw)}c"
            text_widget.tag_add("highlight", pos, end_pos)
            start_idx = end_pos
    text_widget.tag_config("highlight", background="yellow", foreground="black")

def open_search_results(keywords):
    global last_search_time

    start_time = time.time()

    results = search_documents(keywords)

    elapsed = time.time() - start_time
    last_search_time = elapsed

    if not results:
        status_label.config(text="❌ لم يتم العثور على ملفات تحتوي الكلمات المفتاحية.")
        return

    result_window = tk.Toplevel(root)
    result_window.title(f"نتائج البحث عن: {', '.join(keywords)}")
    result_window.geometry("600x500")

    text_widget = tk.Text(result_window, wrap=tk.WORD)
    text_widget.pack(expand=True, fill=tk.BOTH)

    for file_path, content in results:
        text_widget.insert(tk.END, f"📄 {os.path.basename(file_path)}\n")
        preview = content[:1000]
        text_widget.insert(tk.END, preview + "\n\n")

    highlight_text(text_widget, keywords)

    total_files, total_size = get_documents_stats()
    size_mb = total_size / (1024 * 1024)

    status_label.config(text=f"✅ تم العثور على {len(results)} ملف يحتوي الكلمات المفتاحية. "
                             f"عدد المستندات: {total_files}، الحجم الكلي: {size_mb:.2f} ميجابايت، وقت البحث: {elapsed:.2f} ثانية.")

def classify_documents():
    global last_classify_time

    start_time = time.time()

    CLASS_TREE = {
        "Health": ["ambulance", "poisoning", "medications", "emergency", "patient", "blood"],
        "Education": ["curriculum", "university", "students", "school", "education", "lecture"],
        "Computer": ["artificial intelligence", "networks", "computer", "programming", "algorithm", "servers"]
    }

    classified = defaultdict(list)

    for filename in os.listdir(MANUAL_DIR):
        file_path = os.path.join(MANUAL_DIR, filename)
        if not (filename.endswith('.pdf') or filename.endswith('.docx')):
            continue

        text = extract_text(file_path).lower()

        found_category = None
        for category, keywords in CLASS_TREE.items():
            for keyword in keywords:
                if keyword.lower() in text:
                    found_category = category
                    break
            if found_category:
                break

        if found_category:
            classified[found_category].append(filename)
        else:
            classified["غير معروف"].append(filename)

    elapsed = time.time() - start_time
    last_classify_time = elapsed

    result_window = tk.Toplevel(root)
    result_window.title("📂 التصنيف حسب الشجرة")
    result_window.geometry("500x400")

    text_widget = tk.Text(result_window, wrap=tk.WORD)
    text_widget.pack(expand=True, fill=tk.BOTH)

    for category, files in classified.items():
        text_widget.insert(tk.END, f"🔷 {category}:\n")
        for f in files:
            text_widget.insert(tk.END, f"   • {f}\n")
        text_widget.insert(tk.END, "\n")

    total_files, total_size = get_documents_stats()
    size_mb = total_size / (1024 * 1024)

    status_label.config(text=f"✅ تم تصنيف {sum(len(v) for v in classified.values())} ملف. "
                             f"عدد المستندات: {total_files}، الحجم الكلي: {size_mb:.2f} ميجابايت، وقت التصنيف: {elapsed:.2f} ثانية.")

def show_statistics():
    total_files, total_size = get_documents_stats()
    size_mb = total_size / (1024 * 1024)

    stats_window = tk.Toplevel(root)
    stats_window.title("📊 إحصائيات المستندات")
    stats_window.geometry("350x250")

    label1 = tk.Label(stats_window, text=f"عدد المستندات: {total_files}", font=("Arial", 12))
    label1.pack(pady=10)

    label2 = tk.Label(stats_window, text=f"الحجم الكلي: {size_mb:.2f} ميجابايت", font=("Arial", 12))
    label2.pack(pady=10)

    label3 = tk.Label(stats_window, text=f"وقت آخر فرز: {last_sort_time:.2f} ثانية" if last_sort_time else "وقت آخر فرز: لا يوجد", font=("Arial", 12))
    label3.pack(pady=5)

    label4 = tk.Label(stats_window, text=f"وقت آخر بحث: {last_search_time:.2f} ثانية" if last_search_time else "وقت آخر بحث: لا يوجد", font=("Arial", 12))
    label4.pack(pady=5)

    label5 = tk.Label(stats_window, text=f"وقت آخر تصنيف: {last_classify_time:.2f} ثانية" if last_classify_time else "وقت آخر تصنيف: لا يوجد", font=("Arial", 12))
    label5.pack(pady=5)

    close_btn = tk.Button(stats_window, text="إغلاق", command=stats_window.destroy)
    close_btn.pack(pady=15)

# === واجهة المستخدم ===
root = tk.Tk()
root.title("Document Collector")
root.geometry("400x600")

tk.Label(root, text="📁 رفع ملفات PDF/Word يدوياً", font=("Arial", 14)).pack(pady=10)

btn1 = tk.Button(root, text="👡 اختيار ملفات من الجهاز", command=browse_and_save_manual_files, width=30)
btn1.pack(pady=10)

btn_download = tk.Button(root, text="⬇️ تنزيل الملفات من Drive", command=download_files_from_drive, width=30, bg="lightcyan")
btn_download.pack(pady=10)

btn2 = tk.Button(root, text="🔠 فرز الملفات حسب العنوان", command=sort_documents_by_title, width=30, bg="lightblue")
btn2.pack(pady=10)

btn3 = tk.Button(root, text="☁️ رفع الملفات إلى Drive", command=upload_all_documents, width=30, bg="lightgreen")
btn3.pack(pady=10)

btn4 = tk.Button(root, text="🧠 تصنيف المستندات", command=classify_documents, width=30, bg="lightyellow")
btn4.pack(pady=10)

search_entry = tk.Entry(root, width=40)
search_entry.pack(pady=10)

def on_search_clicked():
    query = search_entry.get().strip()
    if query:
        keywords = query.split()
        open_search_results(keywords)
    else:
        status_label.config(text="❌ الرجاء إدخال كلمات للبحث.")

search_btn = tk.Button(root, text="🔍 بحث في المستندات", command=on_search_clicked, bg="orange", width=30)
search_btn.pack(pady=5)

stats_btn = tk.Button(root, text="📊 عرض الإحصائيات", command=show_statistics, width=30, bg="lightgray")
stats_btn.pack(pady=15)

status_label = tk.Label(root, text="", fg="green", font=("Arial", 11))
status_label.pack(pady=5)

root.mainloop()






